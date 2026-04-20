"""
OpenAI-Compatible target adapter.

A single adapter that works with any provider that speaks the OpenAI Chat API:
  - OpenAI (api.openai.com)
  - Groq (api.groq.com) — generous free tier
  - OpenRouter (openrouter.ai) — free models available
  - Together AI (api.together.xyz)
  - Ollama (localhost:11434) — local models, no key needed

Usage:
    target = OpenAICompatibleTarget(provider="groq", api_key="gsk_...")
    target = OpenAICompatibleTarget(provider="openrouter", api_key="sk-or-...")
    target = OpenAICompatibleTarget(provider="ollama")  # no key needed
"""
import os
from typing import Optional, Dict, List
import openai
from vektor.targets.base import BaseTarget


PROVIDER_PRESETS: Dict[str, Dict] = {
    "openai": {
        "base_url": "https://api.openai.com/v1",
        "default_model": "gpt-3.5-turbo",
        "env_key": "OPENAI_API_KEY",
        "pricing": {
            "gpt-3.5-turbo":  {"input": 0.0005,  "output": 0.0015},
            "gpt-4o-mini":    {"input": 0.00015, "output": 0.0006},
            "gpt-4o":         {"input": 0.005,   "output": 0.015},
            "gpt-4":          {"input": 0.03,    "output": 0.06},
        },
    },
    "groq": {
        "base_url": "https://api.groq.com/openai/v1",
        "default_model": "llama-3.1-8b-instant",
        "env_key": "GROQ_API_KEY",
        "pricing": {},  # Free tier — cost tracked as $0
    },
    "openrouter": {
        "base_url": "https://openrouter.ai/api/v1",
        "default_model": "meta-llama/llama-3.1-8b-instruct:free",
        "env_key": "OPENROUTER_API_KEY",
        "pricing": {},  # Free models available on the :free suffix
    },
    "together": {
        "base_url": "https://api.together.xyz/v1",
        "default_model": "meta-llama/Llama-3-8b-chat-hf",
        "env_key": "TOGETHER_API_KEY",
        "pricing": {},
    },
    "ollama": {
        "base_url": "http://localhost:11434/v1",
        "default_model": "llama3",
        "env_key": None,
        "pricing": {},
    },
    # ── Open-source local LLM applications ──────────────────────────────────
    "lmstudio": {
        # LM Studio: https://lmstudio.ai — GUI app, enable Local Server in app
        "base_url": "http://localhost:1234/v1",
        "default_model": "local-model",
        "env_key": None,
        "pricing": {},
    },
    "localai": {
        # LocalAI: https://localai.io — Docker: ghcr.io/mudler/local-ai
        "base_url": "http://localhost:8080/v1",
        "default_model": "gpt-3.5-turbo",
        "env_key": None,
        "pricing": {},
    },
    "openwebui": {
        # Open WebUI: https://github.com/open-webui/open-webui — runs on Ollama
        "base_url": "http://localhost:3000/openai/v1",
        "default_model": "llama3",
        "env_key": "OPENWEBUI_API_KEY",  # Settings > Account > API Keys
        "pricing": {},
    },
    "anythingllm": {
        # AnythingLLM: https://anythingllm.com — all-in-one RAG + chat app
        "base_url": "http://localhost:3001/api/v1/openai",
        "default_model": "@anythingllm/default",
        "env_key": "ANYTHINGLLM_API_KEY",  # Settings > API Keys
        "pricing": {},
    },
    "jan": {
        # Jan: https://jan.ai — Electron desktop app
        "base_url": "http://localhost:1337/v1",
        "default_model": "llama3-8b-instruct",
        "env_key": None,
        "pricing": {},
    },
}

# Universal refusal patterns all OpenAI-compatible models tend to use
UNIVERSAL_REFUSAL_PATTERNS = [
    "i cannot", "i can't", "i will not", "i won't",
    "i must decline", "i am not able to", "i'm not able to",
    "that would be inappropriate", "i should not",
    "it would be wrong", "i am unable to", "i'm unable to",
    "i'm designed to", "my guidelines prevent",
    "as a responsible ai", "as an ai language model",
]

# Per-provider extras (appended on top of universal)
PROVIDER_REFUSAL_EXTRAS: Dict[str, List[str]] = {
    "groq": ["i cannot assist with", "i'm not going to"],
    "openrouter": ["this request goes against", "i'm sorry, but i"],
    "ollama": ["my training prevents me", "i'm not programmed to"],
}


class OpenAICompatibleTarget(BaseTarget):
    """
    Adapter for any OpenAI-API-compatible provider.

    Supports: OpenAI, Groq, OpenRouter, Together AI, Ollama.
    """

    def __init__(
        self,
        provider: str = "openai",
        api_key: Optional[str] = None,
        model: Optional[str] = None,
        base_url: Optional[str] = None,
        system_prompt: Optional[str] = None,
    ):
        preset = PROVIDER_PRESETS.get(provider, {})
        self._provider = provider
        self._pricing = preset.get("pricing", {})

        resolved_model = model or preset.get("default_model", "gpt-3.5-turbo")
        super().__init__(name=f"{provider}-{resolved_model}")
        self.model = resolved_model
        self.system_prompt = system_prompt

        # Resolve API key: explicit > env var > "not-needed" (for Ollama)
        env_key_name = preset.get("env_key")
        resolved_key = api_key or (os.getenv(env_key_name) if env_key_name else None) or "not-needed"

        resolved_url = base_url or preset.get("base_url")
        self.client = openai.OpenAI(api_key=resolved_key, base_url=resolved_url)

    def supports_documents(self) -> bool:
        return True

    def upload_document(self, file_path: str) -> str:
        import os as _os
        if file_path.endswith(".docx"):
            text = self._extract_docx(file_path)
        elif file_path.endswith((".md", ".txt")):
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()
        else:
            raise ValueError(f"Unsupported file type: {file_path}")
        doc_id = _os.path.basename(file_path)
        self.uploaded_documents[doc_id] = text
        return doc_id

    def _extract_docx(self, file_path: str) -> str:
        from docx import Document
        try:
            doc = Document(file_path)
        except Exception as exc:
            raise ValueError(f"Failed to parse DOCX '{file_path}': {exc}") from exc
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text.strip())
        return "\n".join(parts)

    def query(self, prompt: str, stateful: bool = False, use_documents: bool = True, **kwargs) -> str:
        messages = []
        system_content = self._build_system_message(use_documents)
        if system_content:
            messages.append({"role": "system", "content": system_content})
        if stateful and self.conversation_history:
            messages.extend(self.conversation_history)
        messages.append({"role": "user", "content": prompt})

        try:
            response = self.client.chat.completions.create(
                model=self.model,
                messages=messages,
                temperature=kwargs.get("temperature", 0.7),
                max_tokens=kwargs.get("max_tokens", 1000),
            )
            content = response.choices[0].message.content
            if stateful:
                self.conversation_history.append({"role": "user", "content": prompt})
                self.conversation_history.append({"role": "assistant", "content": content})
            self.request_count += 1
            if response.usage:
                self.total_cost += self._calculate_cost(response.usage)
            return content
        except openai.APIError as e:
            raise RuntimeError(f"{self._provider} API error: {self._redact_secrets(str(e))}")

    @staticmethod
    def _redact_secrets(text: str) -> str:
        """Remove API key patterns from error strings before surfacing them."""
        import re as _re
        return _re.sub(r"(sk-[A-Za-z0-9]{6})[A-Za-z0-9]+", r"\1***", text)

    def _build_system_message(self, use_documents: bool) -> Optional[str]:
        parts = []
        if self.system_prompt:
            parts.append(self.system_prompt.strip())
        if use_documents and self.uploaded_documents:
            if parts:
                parts.append("\n\nYou have access to the following documents:")
            else:
                parts.append("You are a helpful assistant. You have access to the following documents:")
            for doc_id, text in self.uploaded_documents.items():
                parts.append(f"\n<document id='{doc_id}'>\n{text}\n</document>")
        return "\n".join(parts) if parts else None

    def estimate_cost(self, prompt: str, response: Optional[str] = None) -> float:
        pricing = self._pricing.get(self.model, {})
        if not pricing:
            return 0.0  # Free tier
        input_tokens = len(prompt) / 4
        output_tokens = 500 if response is None else len(response) / 4
        return (input_tokens / 1000) * pricing["input"] + (output_tokens / 1000) * pricing["output"]

    def _calculate_cost(self, usage) -> float:
        pricing = self._pricing.get(self.model, {})
        if not pricing:
            return 0.0  # Free tier
        return (
            (usage.prompt_tokens / 1000) * pricing["input"] +
            (usage.completion_tokens / 1000) * pricing["output"]
        )

    def get_refusal_patterns(self) -> List[str]:
        patterns = list(UNIVERSAL_REFUSAL_PATTERNS)
        patterns.extend(PROVIDER_REFUSAL_EXTRAS.get(self._provider, []))
        return patterns
