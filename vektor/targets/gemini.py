"""
Google Gemini target adapter (using google.genai SDK).

Free tier: 500 RPD, 10 RPM with Gemini 2.5 Flash.

Usage:
    target = GeminiTarget(api_key="AIza...")
    target = GeminiTarget()  # picks up GEMINI_API_KEY from environment
"""
import os
from typing import Optional, List
from vektor.targets.base import BaseTarget


GEMINI_REFUSAL_PATTERNS = [
    "i cannot", "i can't", "i will not", "i won't",
    "i must decline", "i am not able to", "i'm not able to",
    "that would be inappropriate", "i should not",
    "i am unable to", "i'm unable to",
    "i'm designed to be helpful, but",
    "as a large language model",
    "i'm not able to help with that",
    "i'm programmed to be safe",
    "this goes against my guidelines",
    "i'm designed to avoid",
]


class GeminiTarget(BaseTarget):
    """
    Google Gemini adapter using the current google.genai SDK.

    Default model: gemini-2.5-flash (free tier: 10 RPM, 500 RPD).
    """

    DEFAULT_MODEL = "models/gemini-2.5-flash"

    def __init__(
        self,
        api_key: Optional[str] = None,
        model: str = DEFAULT_MODEL,
        system_prompt: Optional[str] = None,
    ):
        super().__init__(name=f"gemini-{model.split('/')[-1]}")
        self.model = model
        self.system_prompt = system_prompt

        from google import genai
        resolved_key = api_key or os.getenv("GEMINI_API_KEY")
        if not resolved_key:
            raise ValueError(
                "No Gemini API key found. Set GEMINI_API_KEY environment variable "
                "or pass api_key= to GeminiTarget."
            )
        self._client = genai.Client(api_key=resolved_key)

    def supports_documents(self) -> bool:
        return True

    def upload_document(self, file_path: str) -> str:
        import os as _os
        if file_path.endswith(".docx"):
            text = self._extract_docx(file_path)
        elif file_path.endswith((".md", ".txt")):
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()
        else:
            raise ValueError(f"Unsupported file type: {file_path}")
        doc_id = _os.path.basename(file_path)
        self.uploaded_documents[doc_id] = text
        return doc_id

    def _extract_docx(self, file_path: str) -> str:
        from docx import Document
        doc = Document(file_path)
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text.strip())
        return "\n".join(parts)

    def query(self, prompt: str, stateful: bool = False, use_documents: bool = True, **kwargs) -> str:
        from google.genai import types

        system_msg = self._build_system_message(use_documents)
        config = types.GenerateContentConfig(
            system_instruction=system_msg
        ) if system_msg else None

        try:
            if stateful and self.conversation_history:
                # Build conversation contents from history
                contents = []
                for msg in self.conversation_history:
                    role = "user" if msg["role"] == "user" else "model"
                    contents.append(types.Content(role=role, parts=[types.Part(text=msg["content"])]))
                contents.append(types.Content(role="user", parts=[types.Part(text=prompt)]))
                response = self._client.models.generate_content(
                    model=self.model, contents=contents, config=config
                )
            else:
                response = self._client.models.generate_content(
                    model=self.model, contents=prompt, config=config
                )

            content = response.text
            if stateful:
                self.conversation_history.append({"role": "user", "content": prompt})
                self.conversation_history.append({"role": "assistant", "content": content})

            self.request_count += 1
            return content

        except Exception as e:
            raise RuntimeError(f"Gemini API error: {e}")

    def _build_system_message(self, use_documents: bool) -> Optional[str]:
        parts = []
        if self.system_prompt:
            parts.append(self.system_prompt.strip())
        if use_documents and self.uploaded_documents:
            if parts:
                parts.append("\n\nYou have access to the following documents:")
            else:
                parts.append("You are a helpful assistant. You have access to the following documents:")
            for doc_id, text in self.uploaded_documents.items():
                parts.append(f"\n<document id='{doc_id}'>\n{text}\n</document>")
        return "\n".join(parts) if parts else None

    def estimate_cost(self, prompt: str, response: Optional[str] = None) -> float:
        return 0.0  # Free tier

    def get_refusal_patterns(self) -> List[str]:
        return list(GEMINI_REFUSAL_PATTERNS)
